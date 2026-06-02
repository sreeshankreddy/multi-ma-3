"""
Document model for managing document analysis workflow.
Orchestrates different components and manages analysis pipeline.
"""

import os
from typing import Dict, List, Any, Tuple, Optional
from datetime import datetime
import json
import signal
import functools
from utils.pdf_reader import PDFReader
from utils.image_reader import ImageReader
from utils.ocr_engine import OCREngine
from utils.table_extractor import TableExtractor
from utils.text_cleaner import TextCleaner
from utils.summarizer import DocumentSummarizer
from utils.qa_engine import DocumentQAEngine


def timeout(seconds=30):
    """Timeout decorator for long-running operations."""
    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            def handler(signum, frame):
                raise TimeoutError(f"Operation timed out after {seconds} seconds")
            signal.signal(signal.SIGALRM, handler)
            signal.alarm(seconds)
            try:
                result = func(*args, **kwargs)
            finally:
                signal.alarm(0)
            return result
        return wrapper
    return decorator


class DocumentAnalyzer:
    """
    Manages complete document analysis workflow.
    Handles document loading, OCR, table extraction, NLP analysis, and QA.
    """

    def __init__(self):
        """Initialize all analysis components."""
        self.pdf_reader = PDFReader()
        self.image_reader = ImageReader()
        self.ocr_engine = OCREngine()
        self.table_extractor = TableExtractor()
        self.text_cleaner = TextCleaner()
        self.summarizer = DocumentSummarizer()
        self.qa_engine = DocumentQAEngine()

        self.extracted_text = ""
        self.document_metadata = {}
        self.tables = []
        self.analysis_results = {}

    def load_document(self, file_path: str) -> Tuple[bool, str]:
        """
        Load document from file path.

        Args:
            file_path (str): Path to document file.

        Returns:
            Tuple[bool, str]: Success status and extracted text or error message.
        """
        try:
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()

            if ext == '.pdf':
                text, pages_info = self.pdf_reader.extract_text(file_path)
                self.extracted_text = text
                self.document_metadata = {
                    'file_type': 'pdf',
                    'pages': len(pages_info),
                    'page_info': pages_info[:3]  # Store first 3 pages info
                }
                return True, text

            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']:
                image, metadata = self.image_reader.load_image(file_path)
                self.document_metadata = metadata

                # Apply OCR
                ocr_result = self.ocr_engine.extract_text(file_path)
                self.extracted_text = ocr_result['full_text']

                return True, self.extracted_text

            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                self.extracted_text = text
                self.document_metadata = {'file_type': 'txt'}
                return True, text

            elif ext == '.docx':
                from docx import Document
                doc = Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                self.extracted_text = text
                self.document_metadata = {'file_type': 'docx'}
                return True, text

            else:
                return False, f"Unsupported file format: {ext}"

        except Exception as e:
            return False, f"Error loading document: {str(e)}"

    def extract_tables(self, file_path: str) -> Tuple[bool, List[Dict[str, Any]]]:
        """
        Extract tables from PDF document.

        Args:
            file_path (str): Path to PDF file.

        Returns:
            Tuple[bool, List]: Success status and list of tables.
        """
        try:
            _, ext = os.path.splitext(file_path)

            if ext.lower() != '.pdf':
                return False, []

            tables_info = self.table_extractor.extract_tables_with_page_info(file_path)
            self.tables = tables_info

            # Convert DataFrames to dictionaries for JSON serialization
            tables_dict = []
            for table_info in tables_info:
                table_dict = {
                    'table_index': table_info['table_index'],
                    'page_number': table_info['page_number'],
                    'row_count': table_info['row_count'],
                    'column_count': table_info['column_count'],
                    'columns': table_info['columns'],
                    'data': table_info['dataframe'].to_dict('records')
                }
                tables_dict.append(table_dict)

            return True, tables_dict

        except Exception as e:
            return False, []

    def analyze_text(self) -> Dict[str, Any]:
        """
        Perform comprehensive synchronous text analysis on extracted text with robust error handling.

        Returns:
            Dict: Analysis results including summary, keywords, entities, etc.
        """
        if not self.extracted_text:
            return {'error': 'No text to analyze'}

        try:
            results = {}
            
            # Clean text
            try:
                cleaned_text = self.text_cleaner.clean_text(self.extracted_text)
            except Exception as e:
                print(f"Text cleaning error: {e}")
                cleaned_text = self.extracted_text[:10000]  # Fallback: use first 10k chars

            # Generate summary
            try:
                summary = self.summarizer.summarize_text(cleaned_text, max_length=150, min_length=50)
                results['summary'] = summary
            except Exception as e:
                print(f"Summary error: {e}")
                results['summary'] = "Summary generation failed"

            # Extract keywords
            try:
                keywords = self.summarizer.extract_keywords(cleaned_text, num_keywords=15)
                results['keywords'] = keywords
            except Exception as e:
                print(f"Keywords error: {e}")
                results['keywords'] = []

            # Extract entities
            try:
                entities = self.summarizer.extract_entities(cleaned_text)
                results['entities'] = entities
            except Exception as e:
                print(f"Entities error: {e}")
                results['entities'] = {}

            # Extract topics
            try:
                topics = self.summarizer.extract_topics(cleaned_text, num_topics=5)
                results['topics'] = topics
            except Exception as e:
                print(f"Topics error: {e}")
                results['topics'] = []

            # Analyze sentiment
            try:
                sentiment = self.summarizer.analyze_sentiment(cleaned_text)
                results['sentiment'] = sentiment
            except Exception as e:
                print(f"Sentiment error: {e}")
                results['sentiment'] = {"label": "NEUTRAL", "score": 0.0}

            # Generate bullet points
            try:
                bullet_points = self.summarizer.generate_bullet_points(cleaned_text, num_points=5)
                results['bullet_points'] = bullet_points
            except Exception as e:
                print(f"Bullet points error: {e}")
                results['bullet_points'] = []

            # Extract important phrases
            try:
                phrases = self.summarizer.extract_important_phrases(cleaned_text, num_phrases=10)
                results['important_phrases'] = phrases
            except Exception as e:
                print(f"Important phrases error: {e}")
                results['important_phrases'] = []

            # Get document statistics
            try:
                statistics = self.summarizer.get_document_statistics(cleaned_text)
                results['statistics'] = statistics
            except Exception as e:
                print(f"Statistics error: {e}")
                results['statistics'] = {}

            results['analysis_timestamp'] = datetime.now().isoformat()
            self.analysis_results = results
            return results
            
        except Exception as e:
            print(f"Critical analysis error: {e}")
            return {
                'error': f"Error during analysis: {str(e)}",
                'summary': 'Analysis failed',
                'keywords': [],
                'entities': {},
                'topics': [],
                'sentiment': {"label": "UNKNOWN", "score": 0.0},
                'bullet_points': [],
                'important_phrases': [],
                'statistics': {}
            }

    async def analyze_text_async(self) -> Dict[str, Any]:
        """
        Perform comprehensive asynchronous text analysis on extracted text.

        Returns:
            Dict: Analysis results including summary, keywords, entities, etc.
        """
        import asyncio
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self.analyze_text)

    def answer_question(self, question: str) -> Dict[str, Any]:
        """
        Answer a question about the document.

        Args:
            question (str): User question.

        Returns:
            Dict: QA result with answer and confidence.
        """
        try:
            if not self.extracted_text:
                return {'error': 'No document loaded'}

            result = self.qa_engine.contextualized_answer(question, self.extracted_text)
            return result

        except Exception as e:
            return {'error': f"Error answering question: {str(e)}"}

    def batch_qa(self, questions: List[str]) -> List[Dict[str, Any]]:
        """
        Answer multiple questions about the document.

        Args:
            questions (List[str]): List of questions.

        Returns:
            List[Dict]: List of QA results.
        """
        results = []

        for question in questions:
            result = self.answer_question(question)
            results.append(result)

        return results

    def extract_key_information(self) -> Dict[str, Any]:
        """
        Extract key information using predefined questions.

        Returns:
            Dict: Extracted key information.
        """
        try:
            if not self.extracted_text:
                return {'error': 'No document loaded'}

            key_info = self.qa_engine.extract_key_information(self.extracted_text)
            return key_info

        except Exception as e:
            return {'error': f"Error extracting key information: {str(e)}"}

    def get_document_summary(self) -> str:
        """
        Get document summary.

        Returns:
            str: Document summary.
        """
        if self.analysis_results and 'summary' in self.analysis_results:
            return self.analysis_results['summary']
        return ""

    def get_keywords(self) -> List[str]:
        """
        Get extracted keywords.

        Returns:
            List[str]: List of keywords.
        """
        if self.analysis_results and 'keywords' in self.analysis_results:
            return self.analysis_results['keywords']
        return []

    def get_sentiment(self) -> Dict[str, Any]:
        """
        Get sentiment analysis results.

        Returns:
            Dict: Sentiment analysis data.
        """
        if self.analysis_results and 'sentiment' in self.analysis_results:
            return self.analysis_results['sentiment']
        return {}

    def get_tables_as_dataframes(self) -> List:
        """
        Get extracted tables as pandas DataFrames.

        Returns:
            List: List of DataFrames.
        """
        import pandas as pd

        dataframes = []
        for table_info in self.tables:
            df = pd.DataFrame(table_info['data'], columns=table_info['columns'])
            dataframes.append(df)

        return dataframes

    def export_analysis_results(self, output_path: str) -> Tuple[bool, str]:
        """
        Export analysis results to JSON file.

        Args:
            output_path (str): Path to save JSON file.

        Returns:
            Tuple[bool, str]: Success status and message.
        """
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            export_data = {
                'document_metadata': self.document_metadata,
                'analysis_results': self.analysis_results,
                'tables_count': len(self.tables),
                'extracted_text_preview': self.extracted_text[:500]
            }

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)

            return True, f"Results exported to {output_path}"

        except Exception as e:
            return False, f"Error exporting results: {str(e)}"

    def get_analysis_report(self) -> Dict[str, Any]:
        """
        Get comprehensive analysis report.

        Returns:
            Dict: Complete analysis report.
        """
        return {
            'metadata': self.document_metadata,
            'analysis': self.analysis_results,
            'tables': self.tables,
            'text_length': len(self.extracted_text),
            'extraction_timestamp': datetime.now().isoformat()
        }

    def reset(self) -> None:
        """Reset analyzer state."""
        self.extracted_text = ""
        self.document_metadata = {}
        self.tables = []
        self.analysis_results = {}
